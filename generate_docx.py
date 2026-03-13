import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def main():
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Styles
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_normal.paragraph_format.line_spacing = 1.15

    # Heading 1
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Times New Roman'
    font_h1.size = Pt(14)
    font_h1.bold = True
    font_h1.color.rgb = None # Remove blue color
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2
    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Times New Roman'
    font_h2.size = Pt(12)
    font_h2.bold = True
    font_h2.color.rgb = None
    style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_center_para(text, bold=False, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        return p

    # Front Page
    add_center_para("Internship SYNOPSIS (AIML 456)\nOn", bold=True, size=14)
    add_center_para("GMAX Platform\n(B2B Order Management SaaS)", bold=True, size=16)
    add_center_para("Submitted for partial fulfilment of award of the degree of\nBachelor of Technology\nIn\nArtificial Intelligence and Machine Learning", size=12)
    add_center_para("Submitted by\nAditya Tyagi – [Your Enrollment Number]", bold=True, size=12)
    add_center_para("Under the Guidance of\n[Guide Name]\n[Designation]", bold=True, size=12)
    add_center_para("Department of Artificial Intelligence\nDELHI TECHNICAL CAMPUS, GREATER NOIDA\n(Affiliated Guru Gobind Singh Indraprastha University, New Delhi)\nSession 2025-2026 (EVEN SEM)", bold=True, size=12)
    
    doc.add_page_break()

    # Content Pages
    doc.add_heading("Content", level=1)
    
    content_items = [
        "1. Introduction", "2. Problem Statements", "3. Objectives",
        "4. Feasibility Study", "5. Need and Significance", "6. Intended User",
        "7. Abbreviations and Acronyms", "8. Literature Review",
        "9. Proposed Methodology in brief", "10. Hardware Requirements",
        "11. Software Requirements", "12. Diagrams", "13. Gantt Chart", "14. References"
    ]
    for item in content_items:
        doc.add_paragraph(item)
        
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("The GMAX platform is a modern, full-stack B2B Order Management SaaS platform built to revolutionize the wholesale distribution industry. It is specifically designed to replace manual, WhatsApp-based ordering with a highly professional and streamlined digital experience. The platform encapsulates a slick SaaS vibe with an electric aesthetic, incorporating a persistent cart, interactive order tracking, and an analytics dashboard to empower dealers and sales managers. By bridging the gap between traditional ordering methods and modern technological capabilities, GMAX platform aims to greatly reduce administrative overhead and enhance real-time operational visibility.")

    # 2. Problem Statements
    doc.add_heading("2. Problem Statements", level=1)
    doc.add_paragraph("The current state of B2B electrical and wholesale part ordering relies heavily on unorganized communication channels like WhatsApp, generic emails, and phone calls. This leads to severe inefficiencies, order mismatches, manual data entry errors, lack of a centralized tracking mechanism, and minimal analytical insights into sales performance. Dealers struggle to maintain an accurate view of stock, while sales managers lack the real-time data needed to make informed business decisions.")

    # 3. Objectives
    doc.add_heading("3. Objectives", level=1)
    doc.add_paragraph("• To develop a robust, scalable, and responsive B2B platform catering to dealers, sales managers, and administrators.")
    doc.add_paragraph("• To integrate persistent cart functionality, seamless order tracking, and intuitive PDF invoicing.")
    doc.add_paragraph("• To construct an interactive analytics dashboard presenting real-time sales metrics using Recharts.")
    doc.add_paragraph("• To transition users from manual ordering constraints to a seamless digital workflow.")

    # 4. Feasibility Study
    doc.add_heading("4. Feasibility Study", level=1)
    doc.add_paragraph("Technical Feasibility: The implementation employs universally backed, industry-standard technologies such as Next.js 15, Node.js, and MongoDB. These technologies present no major hurdles and boast abundant community support.")
    doc.add_paragraph("Operational Feasibility: The platform mimics a familiar cart-to-checkout layout, making the learning curve extremely low for intended users.")
    doc.add_paragraph("Economic Feasibility: Leveraging cloud services (Render/Vercel) significantly lowers infrastructure and maintenance costs compared to physical servers.")

    # 5. Need and Significance
    doc.add_heading("5. Need and Significance", level=1)
    doc.add_paragraph("In a rapidly digitalizing economy, traditional businesses need sophisticated digital tools to stay competitive. The GMAX platform satisfies the urgent need of wholesale networks requiring transparent, instantaneous order updates, thus eliminating the delay and risk associated with legacy processes. Over time, it will significantly boost productivity and vendor satisfaction.")

    # 6. Intended User
    doc.add_heading("6. Intended User", level=1)
    doc.add_paragraph("• Dealers: Authorized personnel who browse electrical products and place wholesale orders.")
    doc.add_paragraph("• Sales Managers: Personnel responsible for approving workflows, supervising quotes, and managing client relations.")
    doc.add_paragraph("• Admins: Platform administrators managing global catalog data, access roles, and platform settings.")

    # 7. Abbreviations and Acronyms
    doc.add_heading("7. Abbreviations and Acronyms", level=1)
    doc.add_paragraph("SaaS: Software as a Service\nB2B: Business to Business\nJWT: JSON Web Token\nUI/UX: User Interface/User Experience")

    # 8. Literature Review
    doc.add_heading("8. Literature Review", level=1)
    doc.add_paragraph("Traditional wholesale operations have historically been neglected by modern SaaS innovations, relying on ERPs that lack user-centric designs. The shift toward B2B platforms resembles successful consumer-level eCommerce transformations but carries unique challenges, notably bulk negotiations and hierarchy-based approvals. Existing general solutions lack the niche focus required by electrical dealers. Implementing a specialized MERN/Next.js stack addresses performance overheads seen in older tech solutions.")

    # 9. Proposed Methodology in brief
    doc.add_heading("9. Proposed Methodology in brief", level=1)
    doc.add_paragraph("The proposed implementation will follow an Agile development methodology, separated into continuous development sprints:")
    doc.add_paragraph("1. Requirement Analysis & UI/UX Design: Prototyping dark/light electric aesthetic themes.")
    doc.add_paragraph("2. API & Backend Construction: Establishing JWT auth, mongoose schemas for users, products, and analytics.")
    doc.add_paragraph("3. Frontend Integration: Creating Zustand state management, persistent cart, and connecting React components to backend controllers.")
    doc.add_paragraph("4. Testing & Deployment: End-to-end testing, followed by Vercel (Frontend) and Render/Railway (Backend) deployment.")

    # 10. Hardware Requirements
    doc.add_heading("10. Hardware Requirements", level=1)
    doc.add_paragraph("Developer Side:\n• Processor: Intel Core i5/i7, AMD Ryzen 5 or equivalent\n• RAM: 8GB Minimum (16GB Recommended)\n• Storage: 256GB SSD\nClient Side:\n• Any modern computing device (Desktop, Tablet, or Smartphone) capable of running a modern web browser.")

    # 11. Software Requirements
    doc.add_heading("11. Software Requirements", level=1)
    doc.add_heading("Front End", level=2)
    doc.add_paragraph("Framework: Next.js 15, React\nStyling: Tailwind CSS v4\nState Management: Zustand\nCharting: Recharts")
    doc.add_heading("Back End", level=2)
    doc.add_paragraph("Runtime & Framework: Node.js, Express.js\nDatabase: MongoDB with Mongoose ODM\nAuthentication: JSON Web Token (JWT)")

    # 12. Diagrams
    doc.add_heading("12. Diagrams", level=1)
    doc.add_paragraph("The visual mockups emphasize the electric SaaS vibe and functionality requested.")
    
    # Image 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img1_path = r"C:\Users\aditya tyagi\.gemini\antigravity\brain\15ebb73d-fd57-4a4f-948b-aa2050dff5d0\gmax_dashboard_1773214638691.png"
    if os.path.exists(img1_path):
        run = p.add_run()
        run.add_picture(img1_path, width=Inches(6.0))
        p2 = doc.add_paragraph("Figure 1: Gmax Platform Analytics Dashboard Mockup")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Image 2
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img2_path = r"C:\Users\aditya tyagi\.gemini\antigravity\brain\15ebb73d-fd57-4a4f-948b-aa2050dff5d0\gmax_order_tracking_1773214657317.png"
    if os.path.exists(img2_path):
        run3 = p3.add_run()
        run3.add_picture(img2_path, width=Inches(6.0))
        p4 = doc.add_paragraph("Figure 2: Gmax Platform Order Tracking Mockup")
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph("Class, Use Case, DFD (Level 0, 1 & 3), and E-R diagrams will be finalized during the detailed design phase of the project implementation.")

    # 13. Gantt Chart
    doc.add_heading("13. Gantt Chart", level=1)
    doc.add_paragraph("The project is planned over a 16-week period encompassing Requirement Gathering, Design, Backend APIs, Frontend, Integration, Testing, and Final Review.")

    # 14. References
    doc.add_heading("14. References", level=1)
    doc.add_paragraph("[1] React Documentation: https://react.dev/")
    doc.add_paragraph("[2] Next.js Framework Documentation: https://nextjs.org/docs")
    doc.add_paragraph("[3] Node.js and Express Official Guides: https://expressjs.com/")
    doc.add_paragraph("[4] MongoDB Manual: https://www.mongodb.com/docs/manual/")

    doc.save("GMAX_Platform_Synopsis.docx")
    print("Successfully generated GMAX_Platform_Synopsis.docx")

if __name__ == "__main__":
    main()
